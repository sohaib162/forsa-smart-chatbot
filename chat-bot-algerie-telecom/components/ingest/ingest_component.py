"""
IngestComponent — Document Ingestion with MinIO + Vector Indexing
==================================================================
Handles the full document lifecycle:

1. **Upload** — Store raw documents in MinIO S3 (preserving categories)
2. **Parse** — Extract text from PDF, DOCX, ODT, and JSON files
3. **Chunk** — Split text into overlapping chunks with metadata
4. **Index** — Generate embeddings and store in the vector index

This component bridges MinIO (raw document storage) with the
VectorStoreComponent (semantic search index), ensuring every
uploaded document is immediately searchable.

Design
------
- All parsing is done locally (no external APIs).
- PDF: PyMuPDF (fitz) or pdfplumber
- DOCX: python-docx
- JSON: direct load (existing AT document format)
- Chunking uses a sliding window with configurable size/overlap.
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Tuple

from minio import Minio

from components.vector_store.vector_store_component import (
    VectorDocument,
    VectorStoreComponent,
)
from settings import IngestSettings, MinIOSettings

logger = logging.getLogger("forsa.components.ingest")


@dataclass
class IngestResult:
    """Result of a document ingestion operation."""

    filename: str
    s3_key: str
    category: str
    chunks_created: int
    status: str  # "success", "error"
    error: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class IngestComponent:
    """
    Injectable document ingestion component.

    Manages the pipeline: Upload → Parse → Chunk → Index.
    """

    def __init__(
        self,
        ingest_settings: IngestSettings,
        minio_settings: MinIOSettings,
        vector_store: VectorStoreComponent,
        embedding_component: Any,
    ) -> None:
        self._settings = ingest_settings
        self._minio_settings = minio_settings
        self._vector_store = vector_store
        self._embedding = embedding_component
        self._minio_client: Optional[Minio] = None
        self._initialise_minio()

    def _initialise_minio(self) -> None:
        """Connect to MinIO and ensure the bucket exists."""
        logger.info("IngestComponent: Connecting to MinIO at %s …", self._minio_settings.endpoint)
        try:
            self._minio_client = Minio(
                self._minio_settings.endpoint,
                access_key=self._minio_settings.access_key,
                secret_key=self._minio_settings.secret_key,
                secure=self._minio_settings.secure,
            )
            bucket = self._minio_settings.bucket
            if not self._minio_client.bucket_exists(bucket):
                self._minio_client.make_bucket(bucket)
                logger.info("  Created bucket '%s'", bucket)
            else:
                logger.info("  Bucket '%s' exists", bucket)
        except Exception as e:
            logger.warning(
                "IngestComponent: MinIO not available (%s) — "
                "ingestion will work without S3 upload",
                e,
            )
            self._minio_client = None

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    @property
    def is_ready(self) -> bool:
        return self._vector_store.is_ready

    @property
    def minio_available(self) -> bool:
        return self._minio_client is not None

    def ingest_file(
        self,
        filename: str,
        file_data: BinaryIO,
        category: str = "Uncategorized",
        language: str = "FR",
        extra_metadata: Optional[Dict[str, Any]] = None,
    ) -> IngestResult:
        """
        Ingest a single file: upload to MinIO → parse → chunk → index.

        Parameters
        ----------
        filename     : Original filename (e.g. 'convention_P.pdf')
        file_data    : Binary file-like object
        category     : Document category ('Guides', 'Offres', etc.)
        language     : Document language ('FR', 'AR')
        extra_metadata : Additional metadata to attach to each chunk

        Returns
        -------
        IngestResult with status and chunk count
        """
        t0 = time.perf_counter()
        s3_key = f"{category}/{filename}"
        metadata = {
            "filename": filename,
            "category": category,
            "language": language,
            "s3_key": s3_key,
            **(extra_metadata or {}),
        }

        try:
            # Step 1: Upload to MinIO (if available)
            raw_bytes = file_data.read()
            if self._minio_client:
                self._upload_to_minio(s3_key, raw_bytes, filename)

            # Step 2: Parse text from the file
            ext = Path(filename).suffix.lower()
            text = self._parse_file(raw_bytes, ext, filename)

            if not text or len(text.strip()) < 20:
                return IngestResult(
                    filename=filename,
                    s3_key=s3_key,
                    category=category,
                    chunks_created=0,
                    status="error",
                    error="No extractable text found in document",
                )

            # Step 3: Chunk the text
            chunks = self._chunk_text(text, metadata)

            # Step 4: Create VectorDocuments and add to index
            vector_docs = [
                VectorDocument(
                    doc_id=f"{s3_key}::chunk_{i}",
                    text=chunk["text"],
                    metadata=chunk["metadata"],
                )
                for i, chunk in enumerate(chunks)
            ]
            added = self._vector_store.add_documents(vector_docs)

            elapsed = time.perf_counter() - t0
            logger.info(
                "Ingested '%s' | %d chunks | %.1fs",
                filename,
                added,
                elapsed,
            )

            return IngestResult(
                filename=filename,
                s3_key=s3_key,
                category=category,
                chunks_created=added,
                status="success",
                metadata=metadata,
            )

        except Exception as e:
            logger.exception("Failed to ingest '%s'", filename)
            return IngestResult(
                filename=filename,
                s3_key=s3_key,
                category=category,
                chunks_created=0,
                status="error",
                error=str(e),
            )

    def ingest_text(
        self,
        text: str,
        doc_id: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> int:
        """
        Directly ingest raw text (for JSON documents already parsed).

        Returns the number of chunks created.
        """
        meta = metadata or {}
        chunks = self._chunk_text(text, meta)
        vector_docs = [
            VectorDocument(
                doc_id=f"{doc_id}::chunk_{i}",
                text=chunk["text"],
                metadata=chunk["metadata"],
            )
            for i, chunk in enumerate(chunks)
        ]
        return self._vector_store.add_documents(vector_docs)

    def list_ingested_documents(self) -> Dict[str, Any]:
        """Return stats about ingested documents."""
        return {
            "total_chunks": self._vector_store.count(),
            "collection": self._vector_store.collection_name,
            "minio_available": self.minio_available,
        }

    # ------------------------------------------------------------------
    # Internal: MinIO Upload
    # ------------------------------------------------------------------

    def _upload_to_minio(self, s3_key: str, data: bytes, filename: str) -> None:
        """Upload raw file bytes to MinIO."""
        content_type = self._detect_content_type(filename)
        self._minio_client.put_object(
            bucket_name=self._minio_settings.bucket,
            object_name=s3_key,
            data=io.BytesIO(data),
            length=len(data),
            content_type=content_type,
        )
        logger.debug("Uploaded to MinIO: %s (%d bytes)", s3_key, len(data))

    @staticmethod
    def _detect_content_type(filename: str) -> str:
        ext = Path(filename).suffix.lower()
        return {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".odt": "application/vnd.oasis.opendocument.text",
            ".json": "application/json",
            ".txt": "text/plain",
        }.get(ext, "application/octet-stream")

    # ------------------------------------------------------------------
    # Internal: File Parsing
    # ------------------------------------------------------------------

    def _parse_file(self, data: bytes, ext: str, filename: str) -> str:
        """Extract text from a file based on its extension."""
        if ext == ".pdf":
            return self._parse_pdf(data)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(data)
        elif ext == ".json":
            return self._parse_json(data)
        elif ext in (".txt", ".md"):
            return data.decode("utf-8", errors="replace")
        elif ext == ".odt":
            return self._parse_odt(data)
        else:
            logger.warning("Unsupported file type: %s", ext)
            return data.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(data: bytes) -> str:
        """Extract text from PDF using PyMuPDF (fitz) or pdfplumber."""
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=data, filetype="pdf")
            text_parts = []
            for page_num, page in enumerate(doc, 1):
                page_text = page.get_text("text")
                if page_text.strip():
                    text_parts.append(
                        f"[Page {page_num}]\n{page_text}"
                    )
            doc.close()
            return "\n\n".join(text_parts)
        except ImportError:
            pass

        try:
            import pdfplumber

            with pdfplumber.open(io.BytesIO(data)) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(
                            f"[Page {page_num}]\n{page_text}"
                        )
                return "\n\n".join(text_parts)
        except ImportError:
            logger.error(
                "No PDF parser available. Install PyMuPDF or pdfplumber."
            )
            return ""

    @staticmethod
    def _parse_docx(data: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx not installed.")
            return ""

    @staticmethod
    def _parse_json(data: bytes) -> str:
        """Parse JSON documents (existing AT format)."""
        try:
            content = json.loads(data.decode("utf-8"))
            if isinstance(content, list):
                return "\n\n".join(
                    json.dumps(item, ensure_ascii=False) for item in content
                )
            elif isinstance(content, dict):
                return json.dumps(content, ensure_ascii=False, indent=2)
            return str(content)
        except Exception:
            return data.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_odt(data: bytes) -> str:
        """Extract text from ODT files."""
        try:
            from odf.opendocument import load
            from odf import text as odf_text
            from odf.text import P

            doc = load(io.BytesIO(data))
            paragraphs = []
            for p in doc.getElementsByType(P):
                text_content = ""
                for node in p.childNodes:
                    if hasattr(node, "data"):
                        text_content += node.data
                    elif hasattr(node, "__str__"):
                        text_content += str(node)
                if text_content.strip():
                    paragraphs.append(text_content)
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("odfpy not installed.")
            return ""

    # ------------------------------------------------------------------
    # Internal: Text Chunking
    # ------------------------------------------------------------------

    def _chunk_text(
        self,
        text: str,
        base_metadata: Dict[str, Any],
    ) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks with metadata.

        Uses a sliding window approach with page-boundary awareness
        (respects ``[Page N]`` markers from PDF parsing).
        """
        chunk_size = self._settings.chunk_size
        chunk_overlap = self._settings.chunk_overlap
        chunks = []

        # Split by paragraphs first for better semantic boundaries
        paragraphs = re.split(r"\n\n+", text)

        current_chunk = ""
        current_page = 0

        for para in paragraphs:
            # Detect page markers
            page_match = re.match(r"\[Page (\d+)\]", para)
            if page_match:
                current_page = int(page_match.group(1))
                para = re.sub(r"\[Page \d+\]\n?", "", para)

            if len(current_chunk) + len(para) + 1 <= chunk_size:
                current_chunk += ("\n" if current_chunk else "") + para
            else:
                if current_chunk:
                    chunks.append(
                        {
                            "text": current_chunk.strip(),
                            "metadata": {
                                **base_metadata,
                                "page": current_page,
                                "chunk_index": len(chunks),
                            },
                        }
                    )
                # Start new chunk with overlap
                if chunk_overlap > 0 and current_chunk:
                    overlap_text = current_chunk[-chunk_overlap:]
                    current_chunk = overlap_text + "\n" + para
                else:
                    current_chunk = para

        # Don't forget the last chunk
        if current_chunk.strip():
            chunks.append(
                {
                    "text": current_chunk.strip(),
                    "metadata": {
                        **base_metadata,
                        "page": current_page,
                        "chunk_index": len(chunks),
                    },
                }
            )

        return chunks
